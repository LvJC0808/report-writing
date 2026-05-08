#!/usr/bin/env python3
"""Write a DOCX lab report after outline approval, preferably by filling template sections in place."""

from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import dataclass
from pathlib import Path


TOP_HEADING_RE = re.compile(r"^[一二三四五六七八九十]+、")
SECOND_HEADING_RE = re.compile(r"^（[一二三四五六七八九十]+）")
IMAGE_RE = re.compile(r"^!\[(?P<caption>[^\]]*)\]\((?P<path>[^)]+)\)\s*$")

CHINESE_SIZE_PT = {
    "初号": 42,
    "小初": 36,
    "一号": 26,
    "小一": 24,
    "二号": 22,
    "小二": 18,
    "三号": 16,
    "小三": 15,
    "四号": 14,
    "小四": 12,
    "五号": 10.5,
    "小五": 9,
    "六号": 7.5,
    "小六": 6.5,
}


@dataclass
class Block:
    kind: str
    text: str = ""
    path: Path | None = None
    caption: str = ""


@dataclass
class Section:
    heading: str
    blocks: list[Block]


def read_text(path: Path) -> str:
    if not path.exists():
        raise FileNotFoundError(path)
    return path.read_text(encoding="utf-8")


def read_json(path: Path | None) -> dict:
    if path is None:
        return {}
    if not path.exists():
        raise FileNotFoundError(path)
    return json.loads(path.read_text(encoding="utf-8"))


def normalize_heading(text: str) -> str:
    text = re.sub(r"\s+", "", text)
    text = re.sub(r"^[一二三四五六七八九十]+、", "", text)
    text = re.sub(r"^（[一二三四五六七八九十]+）", "", text)
    text = re.sub(r"^\d+[.．、]", "", text)
    text = re.sub(r"^（\d+）", "", text)
    return text.strip()


def is_top_heading(text: str) -> bool:
    return bool(TOP_HEADING_RE.match(text.strip()))


def parse_markdown_sections(text: str, base_dir: Path) -> tuple[str | None, list[Section]]:
    title = None
    sections: list[Section] = []
    current: Section | None = None
    in_code = False
    code_lines: list[str] = []

    def ensure_section() -> Section:
        nonlocal current
        if current is None:
            current = Section("", [])
            sections.append(current)
        return current

    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code and code_lines:
                ensure_section().blocks.append(Block("body", "\n".join(code_lines)))
                code_lines = []
            in_code = not in_code
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not stripped:
            continue
        image_match = IMAGE_RE.match(stripped)
        if image_match:
            image_path = Path(image_match.group("path"))
            if not image_path.is_absolute():
                image_path = base_dir / image_path
            ensure_section().blocks.append(Block("image", path=image_path, caption=image_match.group("caption")))
            continue
        if stripped.startswith("# "):
            title = stripped[2:].strip()
            continue
        if stripped.startswith("## "):
            current = Section(stripped[3:].strip(), [])
            sections.append(current)
            continue
        if stripped.startswith("### "):
            ensure_section().blocks.append(Block("heading2", stripped[4:].strip()))
            continue
        if stripped.startswith("- "):
            ensure_section().blocks.append(Block("body", stripped[2:].strip()))
            continue
        ensure_section().blocks.append(Block("body", stripped))
    return title, sections


def point_size(value) -> float | None:
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return float(value)
    text = str(value).strip()
    if text in CHINESE_SIZE_PT:
        return CHINESE_SIZE_PT[text]
    match = re.search(r"\d+(?:\.\d+)?", text)
    return float(match.group(0)) if match else None


def style_for(requirements: dict, key: str) -> dict:
    aliases = {
        "body": ("body", "正文"),
        "heading1": ("heading1", "heading_1", "一级标题", "level_1"),
        "heading2": ("heading2", "heading_2", "二级标题", "level_2"),
        "caption": ("caption", "题注", "图表题注"),
    }
    for alias in aliases.get(key, (key,)):
        if isinstance(requirements.get(alias), dict):
            return requirements[alias]
    return {}


def apply_run_style(run, style: dict, bold_default: bool = False) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    font = style.get("font") or style.get("font_east_asia") or style.get("中文字体")
    ascii_font = style.get("font_ascii") or style.get("英文字体") or font
    size = point_size(style.get("size_pt") or style.get("size") or style.get("字号"))
    bold = style.get("bold", bold_default)
    if ascii_font:
        run.font.name = ascii_font
    if font:
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), font)
    if size:
        run.font.size = Pt(size)
    run.bold = bool(bold)


def paragraph_after(paragraph, text: str = "", style: str | None = None):
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if text:
        new_paragraph.add_run(text)
    if style:
        try:
            new_paragraph.style = style
        except Exception:
            pass
    return new_paragraph


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    parent.remove(element)


def set_paragraph_text(paragraph, text: str, style: dict, bold_default: bool = False) -> None:
    run = paragraph.add_run(text)
    apply_run_style(run, style, bold_default=bold_default)


def insert_blocks_after(anchor, blocks: list[Block], requirements: dict) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    cursor = anchor
    for block in blocks:
        if block.kind == "heading2":
            cursor = paragraph_after(cursor)
            set_paragraph_text(cursor, block.text, style_for(requirements, "heading2"), bold_default=True)
        elif block.kind == "image":
            if block.path is None or not block.path.exists():
                cursor = paragraph_after(cursor)
                set_paragraph_text(cursor, f"[缺少图片：{block.path}]", style_for(requirements, "body"))
                continue
            cursor = paragraph_after(cursor)
            cursor.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cursor.add_run().add_picture(str(block.path), width=Inches(5.6))
            if block.caption:
                cursor = paragraph_after(cursor)
                cursor.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_paragraph_text(cursor, block.caption, style_for(requirements, "caption"))
        else:
            cursor = paragraph_after(cursor)
            set_paragraph_text(cursor, block.text, style_for(requirements, "body"))


def fill_personal_info(document, personal_info: dict[str, str]) -> None:
    def normalized_with_map(paragraph):
        chars = []
        run_indexes = []
        for run_index, run in enumerate(paragraph.runs):
            for char in run.text:
                if char.isspace() or char == "\u00a0":
                    continue
                chars.append(char)
                run_indexes.append(run_index)
        return "".join(chars), run_indexes

    for label, value in personal_info.items():
        if not value:
            continue
        normalized_label = re.sub(r"[\s\u00a0]+", "", label)
        for paragraph in document.paragraphs:
            combined, run_indexes = normalized_with_map(paragraph)
            label_index = combined.find(normalized_label)
            if label_index == -1:
                continue
            label_end = label_index + len(normalized_label) - 1
            if label_end >= len(run_indexes):
                continue
            last_label_run = run_indexes[label_end]
            blank_runs = [
                run
                for run in paragraph.runs[last_label_run + 1 :]
                if run.text.strip("\u00a0 ").strip() == ""
            ]
            target_run = next((run for run in blank_runs if run.font.underline), None)
            if target_run is None and blank_runs:
                target_run = blank_runs[0]
            if target_run is not None:
                padding = " " * max(1, len(target_run.text) - len(value))
                target_run.text = value + padding
            else:
                paragraph.add_run(f" {value}")
            break


def find_template_sections(document) -> list[tuple[int, str, object]]:
    sections = []
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if is_top_heading(text):
            sections.append((index, normalize_heading(text), paragraph))
    return sections


def fill_sections(document, sections: list[Section], requirements: dict) -> int:
    template_sections = find_template_sections(document)
    content_by_key = {normalize_heading(section.heading): section for section in sections if section.heading}
    matched = 0

    for pos in range(len(template_sections) - 1, -1, -1):
        start_index, key, heading_paragraph = template_sections[pos]
        section = content_by_key.get(key)
        if section is None:
            continue
        end_index = template_sections[pos + 1][0] if pos + 1 < len(template_sections) else len(document.paragraphs)
        for paragraph in reversed(document.paragraphs[start_index + 1 : end_index]):
            delete_paragraph(paragraph)
        insert_blocks_after(heading_paragraph, section.blocks, requirements)
        matched += 1
    return matched


def append_sections(document, title: str | None, sections: list[Section], requirements: dict) -> None:
    if title:
        paragraph = document.add_paragraph()
        set_paragraph_text(paragraph, title, style_for(requirements, "heading1"), bold_default=True)
    for section in sections:
        if section.heading:
            paragraph = document.add_paragraph()
            set_paragraph_text(paragraph, section.heading, style_for(requirements, "heading1"), bold_default=True)
            insert_blocks_after(paragraph, section.blocks, requirements)
        else:
            if document.paragraphs:
                insert_blocks_after(document.paragraphs[-1], section.blocks, requirements)


def write_report(
    template: Path,
    out: Path,
    content_path: Path,
    requirements_path: Path | None,
    personal_info_path: Path | None,
    mode: str,
) -> int:
    if out.exists():
        raise FileExistsError(f"refusing to overwrite existing output: {out}")
    if template.resolve() == out.resolve():
        raise ValueError("output path must differ from template path")

    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise RuntimeError("write_docx_report.py requires python-docx") from exc

    requirements = read_json(requirements_path)
    personal_info = read_json(personal_info_path)
    title, sections = parse_markdown_sections(read_text(content_path), content_path.parent)

    document = Document(template)
    fill_personal_info(document, personal_info)
    matched = 0
    if mode != "append":
        matched = fill_sections(document, sections, requirements)
    if matched == 0:
        if mode == "section-fill":
            raise ValueError("no matching template sections found; rerun with --mode append if acceptable")
        append_sections(document, title, sections, requirements)

    out.parent.mkdir(parents=True, exist_ok=True)
    document.save(out)
    return matched


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--template", type=Path, required=True, help="Template .docx")
    parser.add_argument("--outline", type=Path, required=True, help="Approved report outline markdown")
    parser.add_argument("--content", type=Path, required=True, help="Report content markdown")
    parser.add_argument("--format-requirements", type=Path, help="Confirmed format requirements JSON")
    parser.add_argument("--personal-info", type=Path, help="Confirmed personal-info JSON")
    parser.add_argument("--out", type=Path, required=True, help="Output .docx path")
    parser.add_argument("--mode", choices=("section-fill", "append", "rebuild", "copy-append"), default="section-fill")
    parser.add_argument("--outline-approved", action="store_true", help="Required approval gate")
    args = parser.parse_args(argv)

    if not args.outline_approved:
        print("error: refusing to write DOCX until --outline-approved is supplied", file=sys.stderr)
        return 2
    if args.mode in {"rebuild", "copy-append"}:
        print("warning: legacy mode maps to append fallback", file=sys.stderr)
        args.mode = "append"

    try:
        read_text(args.outline)  # The outline is an approval artifact, not final content.
        matched = write_report(
            template=args.template,
            out=args.out,
            content_path=args.content,
            requirements_path=args.format_requirements,
            personal_info_path=args.personal_info,
            mode=args.mode,
        )
    except Exception as exc:  # noqa: BLE001 - command-line tool should report cleanly.
        print(f"error: {exc}", file=sys.stderr)
        return 1

    print(f"wrote {args.out}")
    print(f"matched_sections={matched}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
